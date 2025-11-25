from docx import Document
doc = Document("test6.docx")
p = doc.paragraphs[4]
p.add_run('문단에 글자 추가').bold = True
doc.save("test7.docx")