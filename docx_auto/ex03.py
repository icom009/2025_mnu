from docx import Document

doc = Document()
doc.add_paragraph("첫번째 문답입니다. 반갑습니다.")

doc.add_paragraph("글머리 문답입니다. 잘 부탁드립니다.", style="ListBullet")
doc.add_paragraph("번호 목록 입니다. 잘 부탁드립니다."             , style="ListNumber")

p = doc.add_paragraph("두번째 문답입니다.")
p.add_run("굵은 글씨입니다. \n").bold = True
p.add_run("기울임 글씨입니다. ").italic = True
p.add_run("밑줄 글씨입니다. ").underline = True

doc.save("test3.docx")