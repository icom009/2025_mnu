from docx import Document
from docx.shared import Pt

doc = Document()
para = doc.add_paragraph("이 글자의 크기를 바꿔볼겁니다.")
para = doc.paragraphs[0].runs
for run in para:
    run.font.size = Pt(24)
    
doc.save("test10.docx")